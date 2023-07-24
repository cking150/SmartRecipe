import tkinter as tk
from tkinter import ttk
import random
import openpyxl
from docx import Document

# Recipe data for the salad, pizza, and ground turkey stir-fry
salad_recipe = {
    "name": "Classic Garden Salad",
    "ingredients": [
        "2 cups mixed salad greens",
        "1/2 cucumber, sliced",
        "1/2 cup cherry tomatoes, halved",
        "1/4 red onion, thinly sliced",
        "1/4 cup sliced black olives",
        "2 tablespoons olive oil",
        "1 tablespoon balsamic vinegar",
        "Salt and pepper to taste",
    ],
    "instructions": [
        "In a large salad bowl, combine salad greens, cucumber, cherry tomatoes, red onion, and black olives.",
        "Drizzle olive oil and balsamic vinegar over the salad.",
        "Season with salt and pepper.",
        "Toss the salad until well combined.",
        "Serve immediately.",
    ],
    "nutrition_facts": {
        "Calories": "120",
        "Carbohydrates": "10g",
        "Protein": "3g",
        "Fat": "8g",
    },
}

pizza_recipe = {
    "name": "Margherita Pizza",
    "ingredients": [
        "1 pre-made pizza crust",
        "1/2 cup tomato sauce",
        "1 cup fresh mozzarella cheese, sliced",
        "1 cup cherry tomatoes, halved",
        "Fresh basil leaves",
        "2 tablespoons olive oil",
        "Salt and pepper to taste",
    ],
    "instructions": [
        "Preheat your oven to the temperature specified on the pizza crust package.",
        "Place the pizza crust on a baking sheet or pizza stone.",
        "Spread the tomato sauce evenly over the crust.",
        "Layer the mozzarella cheese slices and cherry tomato halves on top.",
        "Drizzle olive oil over the pizza and season with salt and pepper.",
        "Bake the pizza in the preheated oven until the cheese is melted and bubbly.",
        "Remove the pizza from the oven and top with fresh basil leaves.",
        "Slice and serve hot.",
    ],
    "nutrition_facts": {
        "Calories": "280",
        "Carbohydrates": "30g",
        "Protein": "14g",
        "Fat": "12g",
    },
}

turkey_stir_fry_recipe = {
    "name": "Ground Turkey Enchilada Stir-Fry",
    "ingredients": [
        "1 lb ground turkey",
        "1 tablespoon olive oil",
        "1 onion, chopped",
        "1 bell pepper, sliced",
        "2 cloves garlic, minced",
        "1 cup corn kernels",
        "1 can (15 oz) black beans, drained and rinsed",
        "1 can (15 oz) enchilada sauce",
        "1 teaspoon chili powder",
        "Salt and pepper to taste",
        "Cooked couscous for serving",
    ],
    "instructions": [
        "In a large skillet, heat olive oil over medium-high heat.",
        "Add chopped onion, sliced bell pepper, and minced garlic. Cook until softened.",
        "Add ground turkey and cook until browned, breaking it up with a spatula.",
        "Stir in corn kernels and black beans.",
        "Pour enchilada sauce over the mixture. Season with chili powder, salt, and pepper.",
        "Simmer the stir-fry for a few minutes until the flavors meld.",
        "Serve the stir-fry over cooked couscous.",
    ],
    "nutrition_facts": {
        "Calories": "380",
        "Carbohydrates": "25g",
        "Protein": "30g",
        "Fat": "18g",
    },
}

# Predefined recipe list with all three recipes
recipe_list = [salad_recipe, pizza_recipe, turkey_stir_fry_recipe]

def generate_recipe():
    recipe = random.choice(recipe_list)
    recipe_text = f"Recipe: {recipe['name']}\n\nIngredients:\n{', '.join(recipe['ingredients'])}\n\nInstructions:\n{', '.join(recipe['instructions'])}\n\nNutrition Facts:\n"
    
    for key, value in recipe['nutrition_facts'].items():
        recipe_text += f"{key}: {value}\n"
    
    result_label.config(text=recipe_text)

def save_to_excel():
    recipe = result_label.cget("text")
    if recipe:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Recipe"
        sheet["A1"] = "Recipe"
        sheet["A2"] = recipe
        filename = "generated_recipe.xlsx"
        workbook.save(filename)
        result_label.config(text=f"Recipe saved to {filename}")

def save_to_word():
    recipe = result_label.cget("text")
    if recipe:
        doc = Document()
        doc.add_heading("Generated Recipe", level=1)
        doc.add_paragraph(recipe)
        filename = "generated_recipe.docx"
        doc.save(filename)
        result_label.config(text=f"Recipe saved to {filename}")

# Tkinter setup
app = tk.Tk()
app.title("Recipe Generator")

# Main frame
frame = ttk.Frame(app, padding="10")
frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

# Recipe generation button
generate_button = ttk.Button(frame, text="Generate Recipe", command=generate_recipe)
generate_button.grid(row=0, column=0, pady=10)

# Save to Excel button
save_excel_button = ttk.Button(frame, text="Save to Excel", command=save_to_excel)
save_excel_button.grid(row=0, column=1, pady=10)

# Save to Word button
save_word_button = ttk.Button(frame, text="Save to Word", command=save_to_word)
save_word_button.grid(row=0, column=2, pady=10)

# Result label to show the generated recipe
result_label = ttk.Label(frame, text="", font=("Helvetica", 12), wraplength=500)
result_label.grid(row=1, column=0, columnspan=3)

# Add padding to all widgets
for child in frame.winfo_children():
    child.grid_configure(padx=5, pady=5)

# Start the Tkinter main loop
app.mainloop()
